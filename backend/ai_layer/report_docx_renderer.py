"""
ai_layer/report_docx_renderer.py – Render báo cáo (HTML fragment) ra file .docx bằng python-docx.

Vì nguồn dữ liệu chính của báo cáo giờ là HTML (sinh trực tiếp từ LLM theo
report_generator_prompt.txt), renderer này PARSE HTML bằng BeautifulSoup rồi
build DOCX tương ứng — không còn đọc JSON blocks như thiết kế trước.

Hỗ trợ các tag HTML mà report_generator_prompt.txt yêu cầu LLM sử dụng:
  p, h2, h3, strong, table/thead/tbody/tr/th/td, ul/ol/li

YÊU CẦU ĐỊNH DẠNG (cố định, không phụ thuộc theme Word):
- TOÀN BỘ văn bản dùng font Times New Roman.
- TOÀN BỘ văn bản màu đen (RGB 0,0,0) — bao gồm cả heading, vì style "Heading 1"/
  "Heading 2" mặc định của Word có màu xanh/xám và font khác Normal, không tự
  kế thừa từ style Normal. Phải set rõ ràng ở mức style VÀ ở mức run (vì một số
  phiên bản Word/python-docx không áp dụng đổi màu cấp style cho run đã tồn tại).
"""

from __future__ import annotations

from io import BytesIO

from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

_BORDER_COLOR = "444444"
_HEADER_FILL = "F1F1F1"
_FONT_NAME = "Times New Roman"
_FONT_COLOR = RGBColor(0, 0, 0)  # đen tuyệt đối

# Các style cần ép font + màu đen. "Normal" áp dụng cho mọi paragraph/list thường,
# "Heading 1".."Heading 4" áp dụng cho <h2>/<h3> (xem mapping trong _render_block),
# "Title" áp dụng cho tiêu đề chính nếu dùng add_heading(level=0).
_STYLES_TO_FORCE = [
    "Normal",
    "Title",
    "Heading 1",
    "Heading 2",
    "Heading 3",
    "Heading 4",
    "List Bullet",
    "List Number",
]


def _force_style_font_black(document: Document, style_name: str) -> None:
    """Ép font + màu đen ở CẤP STYLE, để mọi run mới tạo theo style này kế thừa đúng."""
    try:
        style = document.styles[style_name]
    except KeyError:
        return
    style.font.name = _FONT_NAME
    style.font.color.rgb = _FONT_COLOR
    # Set font cho Đông Á (cần thiết để tiếng Việt có dấu hiển thị đúng font trên 1 số máy Windows)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = r_pr.makeelement(qn("w:rFonts"), {})
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), _FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), _FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), _FONT_NAME)
    r_fonts.set(qn("w:cs"), _FONT_NAME)


def _apply_run_font_black(run) -> None:
    """Ép font + màu đen ở CẤP RUN (phòng trường hợp style không áp dụng được do
    phiên bản Word/python-docx, hoặc run override font cục bộ)."""
    run.font.name = _FONT_NAME
    run.font.color.rgb = _FONT_COLOR
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = r_pr.makeelement(qn("w:rFonts"), {})
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), _FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), _FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), _FONT_NAME)
    r_fonts.set(qn("w:cs"), _FONT_NAME)


def _force_all_runs_in_paragraph(paragraph) -> None:
    """Handle the all runs in paragraph."""

    for run in paragraph.runs:
        _apply_run_font_black(run)


def _is_centered(tag: Tag) -> bool:
    """Determine whether the centered."""

    style = tag.get("style", "") or ""
    return "text-align:center" in style.replace(" ", "")


def _set_cell_border(cell) -> None:
    """Handle the cell border."""

    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.makeelement(qn("w:tcBorders"), {})
    for edge in ("top", "left", "bottom", "right"):
        element = tc_pr.makeelement(
            qn(f"w:{edge}"),
            {
                qn("w:val"): "single",
                qn("w:sz"): "4",
                qn("w:space"): "0",
                qn("w:color"): _BORDER_COLOR,
            },
        )
        borders.append(element)
    tc_pr.append(borders)


def _set_cell_shading(cell, fill: str) -> None:
    """Handle the cell shading."""

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(
        qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): fill}
    )
    tc_pr.append(shd)


def _add_paragraph_with_inline_formatting(
    document: Document, tag: Tag, style: str | None = None
):
    """Thêm 1 paragraph, giữ định dạng <strong>/<b> bên trong (bold từng run riêng),
    luôn ép font Times New Roman + màu đen cho mọi run."""
    paragraph = document.add_paragraph(style=style)
    if _is_centered(tag):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if not tag.contents:
        run = paragraph.add_run(tag.get_text())
        _apply_run_font_black(run)
        return paragraph

    for child in tag.children:
        if isinstance(child, Tag) and child.name in ("strong", "b"):
            run = paragraph.add_run(child.get_text())
            run.bold = True
            _apply_run_font_black(run)
        elif isinstance(child, Tag) and child.name in ("em", "i"):
            run = paragraph.add_run(child.get_text())
            run.italic = True
            _apply_run_font_black(run)
        elif isinstance(child, Tag) and child.name == "br":
            run = paragraph.add_run()
            run.add_break()
        else:
            text = child.get_text() if isinstance(child, Tag) else str(child)
            if text.strip():
                run = paragraph.add_run(text)
                _apply_run_font_black(run)
    return paragraph


def _render_table(document: Document, table_tag: Tag) -> None:
    """Render the table."""

    rows_data: list[list[str]] = []
    header_cells: list[str] = []

    thead = table_tag.find("thead")
    if thead:
        header_row = thead.find("tr")
        if header_row:
            header_cells = [
                cell.get_text(strip=True) for cell in header_row.find_all(["th", "td"])
            ]

    tbody = table_tag.find("tbody") or table_tag
    for tr in tbody.find_all("tr"):
        cells = [cell.get_text(strip=True) for cell in tr.find_all(["td", "th"])]
        if cells and cells != header_cells:
            rows_data.append(cells)

    if not header_cells and not rows_data:
        return

    col_count = max(len(header_cells), max((len(r) for r in rows_data), default=0), 1)
    table = document.add_table(rows=0, cols=col_count)
    table.autofit = False
    col_width = Inches(6.3 / col_count)
    for col in table.columns:
        col.width = col_width

    def _fill_cell(cell, text: str, bold: bool = False) -> None:
        """Handle the cell."""

        cell.text = text
        cell.width = col_width
        _set_cell_border(cell)
        for paragraph in cell.paragraphs:
            if not paragraph.runs:
                # cell.text = "" có thể không tạo run nào; đảm bảo vẫn có 1 run để áp font
                paragraph.add_run("")
            for run in paragraph.runs:
                run.bold = bold
                _apply_run_font_black(run)

    if header_cells:
        header_row_obj = table.add_row()
        for index in range(col_count):
            cell = header_row_obj.cells[index]
            _fill_cell(
                cell,
                header_cells[index] if index < len(header_cells) else "",
                bold=True,
            )
            _set_cell_shading(cell, _HEADER_FILL)

    for row in rows_data:
        data_row = table.add_row()
        for index in range(col_count):
            cell = data_row.cells[index]
            _fill_cell(cell, row[index] if index < len(row) else "")

    document.add_paragraph()


def render_report_docx(markdown_content: str) -> bytes:
    """Parse Markdown báo cáo, convert sang HTML và render ra bytes của file .docx."""
    import markdown

    html = markdown.markdown(markdown_content or "", extensions=["tables"])
    wrapped_html = f'<div class="report">{html}</div>'
    soup = BeautifulSoup(wrapped_html, "html.parser")

    document = Document()

    section = document.sections[0]
    section.page_width = Inches(8.27)  # A4
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(0.9)

    # Ép font Times New Roman + màu đen ở CẤP STYLE trước khi thêm bất kỳ nội dung nào,
    # để mọi paragraph/heading tạo ra sau đó kế thừa đúng ngay từ đầu.
    for style_name in _STYLES_TO_FORCE:
        _force_style_font_black(document, style_name)
    document.styles["Normal"].font.size = Pt(13)

    # Lấy các node cấp cao nhất. Hỗ trợ trường hợp LLM bọc toàn bộ trong <div class="report">
    container = soup.find("div", class_="report")
    if not container:
        container = soup.body if soup.body else soup

    for tag in container.find_all(
        ["p", "h2", "h3", "table", "ul", "ol"], recursive=False
    ):
        if tag.name == "p":
            paragraph = _add_paragraph_with_inline_formatting(document, tag)
            if "CỘNG HÒA XÃ HỘI" in tag.get_text():
                # Tiêu ngữ hành chính dùng cỡ chữ 14
                for run in paragraph.runs:
                    run.font.size = Pt(14)
        elif tag.name == "h2":
            paragraph = document.add_heading(tag.get_text(strip=True), level=1)
            if _is_centered(tag):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if "BÁO CÁO GIAO BAN" in tag.get_text():
                for run in paragraph.runs:
                    run.font.size = Pt(16)
                    run.bold = True
            _force_all_runs_in_paragraph(paragraph)
        elif tag.name == "h3":
            paragraph = document.add_heading(tag.get_text(strip=True), level=2)
            _force_all_runs_in_paragraph(paragraph)
        elif tag.name == "table":
            _render_table(document, tag)
        elif tag.name in ("ul", "ol"):
            style_name = "List Number" if tag.name == "ol" else "List Bullet"
            for li in tag.find_all("li", recursive=False):
                paragraph = document.add_paragraph(
                    li.get_text(strip=True), style=style_name
                )
                _force_all_runs_in_paragraph(paragraph)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
