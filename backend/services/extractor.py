"""
services/extractor.py – Trích xuất nội dung văn bản từ các loại file.

Hỗ trợ:
  - PDF      → pdfplumber (đọc text + tables) hoặc fallback pypdf/docling
  - Word     → python-docx (paragraphs + tables)
  - Excel    → openpyxl (cell values của tất cả sheets)
  - Image    → base64 encode để gửi vision model
  - Fallback → đọc raw bytes, cố gắng decode UTF-8

Mỗi hàm trả về dict:
  {
    "text": str,        # nội dung văn bản (có thể rỗng với ảnh)
    "image_b64": str,   # base64 data URI (chỉ có với ảnh)
    "page_count": int,
    "is_image": bool,
    "error": str,       # rỗng nếu không lỗi
  }
"""

from __future__ import annotations

import base64
import io
import traceback
from pathlib import Path
from typing import Any


ExtractionResult = dict[str, Any]

_EMPTY: ExtractionResult = {
    "text": "",
    "image_b64": "",
    "page_count": 0,
    "is_image": False,
    "error": "",
}


def _result(**kwargs) -> ExtractionResult:
    return {**_EMPTY, **kwargs}


# ══════════════════════════════════════════════════════════════════════
# PDF
# ══════════════════════════════════════════════════════════════════════

def extract_pdf(file_path: Path) -> ExtractionResult:
    """Trích xuất text và bảng từ PDF bằng pdfplumber. Fallback sang pypdf."""
    try:
        import pdfplumber  # noqa: PLC0415

        parts: list[str] = []
        page_count = 0

        with pdfplumber.open(str(file_path)) as pdf:
            page_count = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    parts.append(f"--- Trang {i + 1} ---\n{page_text}")

                # Trích xuất bảng (nếu có)
                tables = page.extract_tables()
                for tbl_idx, table in enumerate(tables):
                    rows = []
                    for row in table:
                        cleaned = [str(cell or "").strip() for cell in row]
                        rows.append(" | ".join(cleaned))
                    if rows:
                        parts.append(
                            f"[Bảng {tbl_idx + 1} – Trang {i + 1}]\n"
                            + "\n".join(rows)
                        )

        text = "\n\n".join(parts)
        if not text.strip():
            text = "[PDF không có text layer – có thể là ảnh scan]"

        return _result(text=text, page_count=page_count)

    except ImportError:
        # Fallback to pypdf if pdfplumber is not installed
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(file_path))
            page_count = len(reader.pages)
            parts = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    parts.append(f"--- Trang {i + 1} ---\n{page_text}")
            text = "\n\n".join(parts)
            if not text.strip():
                text = "[PDF không có text layer – có thể là ảnh scan]"
            return _result(text=text, page_count=page_count)
        except Exception as exc:
            return _result(error=f"Lỗi đọc PDF fallback: {exc}")
    except Exception as exc:  # noqa: BLE001
        return _result(error=f"Lỗi đọc PDF: {exc}\n{traceback.format_exc()}")


# ══════════════════════════════════════════════════════════════════════
# Word (.docx)
# ══════════════════════════════════════════════════════════════════════

def extract_word(file_path: Path) -> ExtractionResult:
    """Trích xuất text từ .docx (paragraphs + tables)."""
    try:
        from docx import Document  # noqa: PLC0415

        doc = Document(str(file_path))
        parts: list[str] = []

        # Paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Tables
        for tbl_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append(f"[Bảng {tbl_idx + 1}]\n" + "\n".join(rows))

        text = "\n".join(parts)
        if not text.strip():
            text = "[File Word không có nội dung văn bản]"

        return _result(text=text, page_count=1)

    except ImportError:
        return _result(error="python-docx chưa được cài. Chạy: pip install python-docx")
    except Exception as exc:  # noqa: BLE001
        return _result(error=f"Lỗi đọc Word: {exc}\n{traceback.format_exc()}")


# ══════════════════════════════════════════════════════════════════════
# Excel (.xlsx / .xls)
# ══════════════════════════════════════════════════════════════════════

def extract_excel(file_path: Path) -> ExtractionResult:
    """Trích xuất dữ liệu từ tất cả sheet của file Excel."""
    try:
        import openpyxl  # noqa: PLC0415

        wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
        parts: list[str] = []
        total_rows = 0

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_text: list[str] = []
            for row in ws.iter_rows(values_only=True):
                # Bỏ qua dòng toàn None
                if not any(cell is not None for cell in row):
                    continue
                cells = [str(c) if c is not None else "" for c in row]
                rows_text.append(" | ".join(cells))
                total_rows += 1
                # Giới hạn 500 dòng/sheet để tránh prompt quá dài
                if total_rows > 500:
                    rows_text.append("[... còn nhiều dòng không hiển thị ...]")
                    break

            if rows_text:
                parts.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows_text))

        wb.close()
        text = "\n\n".join(parts)
        if not text.strip():
            text = "[File Excel không có dữ liệu]"

        return _result(text=text, page_count=len(wb.sheetnames))

    except ImportError:
        return _result(error="openpyxl chưa được cài. Chạy: pip install openpyxl")
    except Exception as exc:  # noqa: BLE001
        return _result(error=f"Lỗi đọc Excel: {exc}\n{traceback.format_exc()}")


# ══════════════════════════════════════════════════════════════════════
# Image
# ══════════════════════════════════════════════════════════════════════

def extract_image(file_path: Path) -> ExtractionResult:
    """
    Encode ảnh thành base64 data URI để gửi vision model.
    Tự động resize nếu ảnh quá lớn (> 2000px) để tiết kiệm tokens.
    """
    try:
        from PIL import Image as PILImage  # noqa: PLC0415

        MAX_DIM = 2000  # pixel

        with PILImage.open(str(file_path)) as img:
            # Chuyển sang RGB nếu cần (bỏ alpha channel)
            if img.mode not in ("RGB", "L"):
                img = img.convert("RGB")

            # Resize nếu quá lớn
            w, h = img.size
            if max(w, h) > MAX_DIM:
                ratio = MAX_DIM / max(w, h)
                new_size = (int(w * ratio), int(h * ratio))
                img = img.resize(new_size, PILImage.LANCZOS)

            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=85)
            buf.seek(0)
            b64 = base64.b64encode(buf.read()).decode("utf-8")

        mime = "image/jpeg"
        data_uri = f"data:{mime};base64,{b64}"

        return _result(
            text="",  # AI sẽ đọc trực tiếp từ ảnh
            image_b64=data_uri,
            page_count=1,
            is_image=True,
        )

    except ImportError:
        # Fallback if Pillow is not installed: read raw bytes and base64 encode
        try:
            with open(file_path, "rb") as f:
                b64 = base64.b64encode(f.read()).decode("utf-8")
            mime = "image/jpeg"
            suffix = file_path.suffix.lower()
            if suffix in {".png", ".gif", ".webp", ".bmp"}:
                mime = f"image/{suffix[1:]}"
            data_uri = f"data:{mime};base64,{b64}"
            return _result(
                text="",
                image_b64=data_uri,
                page_count=1,
                is_image=True,
            )
        except Exception as exc:
            return _result(error=f"Lỗi đọc ảnh fallback: {exc}")
    except Exception as exc:  # noqa: BLE001
        return _result(error=f"Lỗi đọc ảnh: {exc}\n{traceback.format_exc()}")


# ══════════════════════════════════════════════════════════════════════
# Dispatcher
# ══════════════════════════════════════════════════════════════════════

def extract(file_path: Path, file_type: str) -> ExtractionResult:
    """
    Entry point chính. Tự chọn hàm trích xuất theo loại file.

    Args:
        file_path: Đường dẫn tuyệt đối đến file.
        file_type: Giá trị FileType enum ('pdf', 'word', 'excel', 'image').

    Returns:
        ExtractionResult dict.
    """
    if not file_path.exists():
        return _result(error=f"File không tồn tại: {file_path}")

    dispatch = {
        "pdf":   extract_pdf,
        "word":  extract_word,
        "excel": extract_excel,
        "image": extract_image,
    }

    fn = dispatch.get(file_type)
    if fn is None:
        # Thử đọc raw text
        try:
            text = file_path.read_text(encoding="utf-8", errors="ignore")
            return _result(text=text[:8000], page_count=1)
        except Exception as exc:  # noqa: BLE001
            return _result(error=f"Loại file không được hỗ trợ: {file_type}. {exc}")

    return fn(file_path)
