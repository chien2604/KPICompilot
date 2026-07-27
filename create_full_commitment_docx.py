import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = docx.Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)

# Title
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc\n\n")
run.bold = True

title = doc.add_paragraph('CAM KẾT QUYỀN SỞ HỮU TRÍ TUỆ\nSẢN PHẨM/ GIẢI PHÁP DỰ THI\nCUỘC THI SÁNG KIẾN KHOA HỌC 2025')
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title.runs[0].bold = True
title.runs[0].font.size = Pt(14)

# Section I
doc.add_paragraph().add_run('I. ĐƠN VỊ THỰC HIỆN').bold = True
doc.add_paragraph('Tên đơn vị: …………………………………………………………………………………')
doc.add_paragraph('Người đại diện: ……………………………………………………………………………')
doc.add_paragraph('Giới tính:           ☐ Nam              ☐ Nữ')
doc.add_paragraph('Năm sinh: …………………………………………………………………………………')
doc.add_paragraph('CCCD: ……………………………………………………………………………………')
doc.add_paragraph('Địa chỉ liên hệ: ……………………………………………………………………………')
doc.add_paragraph('Số điện thoại: ……………………………………………………………………………')
doc.add_paragraph('Địa chỉ Email: ……………………………………………………………………………')
doc.add_paragraph('Địa chỉ email của Người giám hộ (đối với tác giả/nhóm tác giả dưới 18 tuổi):\n…………………………………………………………………………………………………')
doc.add_paragraph('Các thành viên nhóm:\n…………………………………………………………………………………………………\n…………………………………………………………………………………………………\n…………………………………………………………………………………………………')

# Section II
doc.add_paragraph().add_run('II. GIẢI PHÁP/SẢN PHẨM DỰ THI:').bold = True
doc.add_paragraph().add_run('1. Tên giải pháp/sản phẩm/nền tảng:').bold = True
doc.add_paragraph('"AI KPI Copilot for Government" - Trợ lý ảo AI hỗ trợ đánh giá, quản lý KPI và công việc dành cho Cơ quan Nhà nước.')

doc.add_paragraph().add_run('2. Chi tiết bản quyền của sản phẩm:').bold = True
doc.add_paragraph('(Nêu rõ các tính năng hay thành phần nào của sản phẩm dự thi là hoàn toàn do tác giả phát triển hay xây dựng; nêu rõ các tính năng hay thành phần nào của sản phẩm dự thi được sử dụng lại hay được phát triển trên một phần hay toàn bộ sản phẩm nào khác đã có sẵn, chỉ rõ ở đâu có thể tìm kiếm thông tin về sản phẩm có sẵn này):').italic = True

doc.add_paragraph('a) Các thành phần do nhóm tác giả hoàn toàn tự nghiên cứu, thiết kế và phát triển:\n'
'- Kiến trúc hệ thống tổng thể: Thiết kế luồng nghiệp vụ và mô hình dữ liệu bám sát theo quy định về đánh giá, xếp loại chất lượng cán bộ, công chức, viên chức của Nhà nước.\n'
'- Khối Logic nghiệp vụ (Rule Engine): Bộ quy tắc chấm điểm và đánh giá KPI được xây dựng riêng, có khả năng tự động tổng hợp, đối chiếu kết quả thực hiện nhiệm vụ với mục tiêu đề ra.\n'
'- Hệ thống Trợ lý ảo (AI Copilot): Tự thiết kế hệ thống kịch bản (prompts) và luồng xử lý riêng để Trợ lý ảo hiểu và trả lời chính xác các nghiệp vụ, thuật ngữ hành chính công bằng tiếng Việt.\n'
'- Hệ thống Phân tích minh chứng: Thuật toán tự động bóc tách thông tin từ các tài liệu, báo cáo (văn bản, PDF) do chuyên viên nộp lên để hỗ trợ công tác thẩm định tự động.\n'
'- Bản đồ nhiệt (Risk Heatmap): Thuật toán phân tích mối liên kết giữa các phòng ban, cá nhân và đầu việc để phát hiện, cảnh báo sớm các điểm nghẽn, rủi ro chậm trễ.\n'
'- Giao diện người dùng: Thiết kế Bảng điều khiển (Dashboard) quản trị trực quan, thân thiện, phù hợp với thói quen sử dụng của cán bộ nhà nước.\n\n'
'b) Các thành phần kế thừa, sử dụng lại từ các nền tảng mã nguồn mở (có sẵn):\n'
'- Lõi Xử lý Ngôn ngữ tự nhiên (LLM): Kế thừa các Mô hình ngôn ngữ lớn tiêu chuẩn (mã nguồn mở) để làm nền tảng trí tuệ nhân tạo cốt lõi cho việc phân tích văn bản và sinh báo cáo tự động.\n'
'- Khối Lưu trữ dữ liệu lớn: Kế thừa các hệ quản trị cơ sở dữ liệu quan hệ, cơ sở dữ liệu véc-tơ (Vector DB) và đồ thị tri thức (Knowledge Graph) mã nguồn mở để lưu trữ an toàn thông tin nghiệp vụ và tối ưu hóa tìm kiếm.\n'
'- Công cụ lập trình: Ứng dụng được phát triển dựa trên các nền tảng (framework) lập trình web tiêu chuẩn, mã nguồn mở, phổ biến và được cộng đồng quốc tế công nhận.')

doc.add_paragraph().add_run('3. Danh sách các tài liệu tham khảo chứng minh bản quyền của sản phẩm').bold = True
doc.add_paragraph('(Sản phẩm, giải pháp đã được đăng trên các tạp chí, bài báo khoa học; Sản phẩm demo; Bài phát biểu chia sẻ ý tưởng công khai,…)').italic = True
doc.add_paragraph('- Kho mã nguồn (Source code) toàn bộ hệ thống do nhóm tác giả tự lập trình và lưu trữ.\n'
'- Tài liệu thuyết minh giải pháp kỹ thuật, kiến trúc hệ thống đính kèm trong hồ sơ dự thi.\n'
'- Phiên bản Demo chạy trực tuyến chứng minh đầy đủ các chức năng thực tế của sản phẩm.')

# Section III
doc.add_paragraph().add_run('III. CAM KẾT').bold = True
doc.add_paragraph('Đơn vị thực hiện (người sản xuất hay nhóm sản xuất) cam kết toàn bộ thông tin về quyền sở hữu trí tuệ đối với sản phẩm, giải pháp dự thi “Cuộc thi Sáng kiến Khoa học 2025” do Báo VnExpress tổ chức được cung cấp trên đây là hoàn toàn đúng sự thật.')
doc.add_paragraph('Đơn vị thực hiện bảo đảm ý tưởng, sản phẩm dự thi hoàn toàn do cá nhân/nhóm đăng ký thực hiện, thiết kế, phát triển và chưa từng đạt giải tại các cuộc thi khác.')
doc.add_paragraph('Trong trường hợp có tranh chấp về bản quyền sở hữu đối với một phần hay toàn bộ sản phẩm dự thi, hoặc phát hiện vi phạm về tiêu chí giải thưởng, Đơn vị thực hiện (người sản xuất hay nhóm sản xuất) hoàn toàn chịu trách nhiệm trước pháp luật và Ban tổ chức Cuộc thi về tính trung thực của bản Cam kết này.')

# Footer signatures
table = doc.add_table(rows=1, cols=2)
table.autofit = False
table.columns[0].width = Inches(3.5)
table.columns[1].width = Inches(3.5)

cell_0 = table.cell(0, 0)
p_0 = cell_0.paragraphs[0]
p_0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p_0.add_run('Xác nhận của Người giám hộ\n').bold = True
p_0.add_run('(Đối với tác giả/nhóm tác giả dưới 18 tuổi)')

cell_1 = table.cell(0, 1)
p_1 = cell_1.paragraphs[0]
p_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p_1.add_run('Bản Cam kết khai tại: ..........................\nNgày …… tháng …… năm ……\n')
p_1.add_run('Họ tên, chữ ký của tác giả\n\n\n\n(Hoặc đại diện Nhóm tác giả)').bold = True

doc.add_paragraph()
doc.add_paragraph('* Đối với Nhóm tác giả khuyến khích bổ sung Họ tên và chữ ký của các Đồng tác giả.').italic = True

doc.save('Ban_Cam_Ket_Hoan_Chinh.docx')
