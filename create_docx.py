import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = docx.Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)

# Title
title = doc.add_paragraph('BẢN MÔ TẢ Ý TƯỞNG, SẢN PHẨM')
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title.runs[0].bold = True
title.runs[0].font.size = Pt(16)

# Content sections
sections = [
    ("1. Tên ý tưởng, sản phẩm:", 
     "\"AI KPI Copilot for Government\" - Trợ lý ảo AI tiên phong trong việc hỗ trợ đánh giá, quản lý KPI và tối ưu hóa quy trình thực thi công việc dành riêng cho hệ thống Cơ quan Hành chính Nhà nước."),
    
    ("2. Vấn đề cần giải quyết:", 
     "Trong bối cảnh chuyển đổi số quốc gia đang diễn ra mạnh mẽ, việc quản lý hiệu suất công việc (KPI) tại các Sở, Ban, Ngành và cơ quan quản lý nhà nước vẫn còn tồn tại nhiều bất cập:\n"
     "- Tính thủ công và rườm rà trong đánh giá: Việc lập báo cáo, nộp minh chứng và đánh giá KPI cuối tháng/quý thường được thực hiện qua nhiều biểu mẫu Excel, văn bản giấy hoặc các hệ thống phân tán. Điều này làm mất rất nhiều thời gian của chuyên viên và cán bộ quản lý.\n"
     "- Khó khăn trong xác thực minh chứng: Khi một cán bộ nộp tài liệu minh chứng cho công việc đã hoàn thành, người quản lý phải tự đọc và đối chiếu từng trang tài liệu để xác nhận mức độ hoàn thành. Khâu này dễ sinh ra sai sót, tốn kém thời gian và đôi khi mang nặng tính cảm tính.\n"
     "- Thiếu tầm nhìn tổng thể và cảnh báo rủi ro (Risk Management): Lãnh đạo cấp cao (Giám đốc, Phó giám đốc Sở) không có một bức tranh toàn cảnh (dashboard) theo thời gian thực về tiến độ công việc của toàn bộ các phòng ban. Các điểm nghẽn, sự cố chậm trễ thường chỉ được phát hiện khi đã quá hạn hoặc qua các kỳ họp giao ban định kỳ, dẫn đến sự bị động trong điều hành.\n"
     "- Tra cứu và tổng hợp thông tin khó khăn: Lãnh đạo cần thông tin báo cáo nhanh gọn để ra quyết định nhưng phải đợi các phòng ban tổng hợp mất nhiều ngày. Thiếu một công cụ tương tác trực tiếp bằng ngôn ngữ tự nhiên để truy vấn dữ liệu tức thời."),
    
    ("3. Giải pháp ứng dụng trí tuệ nhân tạo dự kiến:", 
     "\"AI KPI Copilot for Government\" là một giải pháp toàn diện (end-to-end) giải quyết triệt để các bài toán trên thông qua việc ứng dụng các mô hình Trí tuệ nhân tạo (AI) tiên tiến:\n"
     "- Trợ lý ảo thông minh (AI Copilot) dành cho Lãnh đạo: Khác với các hệ thống phần mềm truyền thống (chỉ có nút bấm và biểu đồ), hệ thống cung cấp một chatbot thông minh, cho phép Lãnh đạo giao tiếp bằng ngôn ngữ tự nhiên (Tiếng Việt). Ví dụ: Lãnh đạo có thể hỏi \"Tiến độ giải ngân vốn đầu tư công của phòng Kế hoạch tháng này ra sao?\", AI sẽ ngay lập tức tổng hợp số liệu và trả lời.\n"
     "- Tự động hóa Phân tích minh chứng (Document AI & Vector Search): Khi chuyên viên tải lên các báo cáo, tài liệu minh chứng, hệ thống sử dụng công nghệ OCR (đối với ảnh/PDF) kết hợp AI đọc hiểu ngôn ngữ tự nhiên. Các tài liệu được nhúng (embedding) thành vector và lưu trữ bằng pgvector để dễ dàng tìm kiếm theo ngữ nghĩa. AI sẽ đối chiếu nội dung tài liệu với yêu cầu của nhiệm vụ để tự động đề xuất tỷ lệ hoàn thành (0-100%).\n"
     "- Bản đồ nhiệt cảnh báo rủi ro (Risk Heatmap) dựa trên Đồ thị Tri thức (Knowledge Graph): Hệ thống sử dụng KùzuDB để xây dựng mạng lưới đồ thị liên kết giữa mục tiêu (KPI) - phòng ban - cá nhân - nhiệm vụ. Thuật toán AI sẽ phân tích đồ thị này để tìm ra các \"nút thắt cổ chai\" (bottlenecks) đang làm chậm tiến độ chung, sau đó hiển thị trực quan dưới dạng Bản đồ nhiệt (Heatmap) với các cấp độ màu sắc cảnh báo rủi ro.\n"
     "- Tự động sinh báo cáo giao ban (Generative AI): Khai thác sức mạnh của Mô hình Ngôn ngữ Lớn (LLMs), hệ thống có khả năng tự động tổng hợp toàn bộ dữ liệu hoạt động trong tuần/tháng để sinh ra bản báo cáo giao ban hoàn chỉnh, chỉn chu, giúp thư ký hoặc chánh văn phòng tiết kiệm 90% thời gian chuẩn bị."),
    
    ("4. Công nghệ, phương thức triển khai:", 
     "Nền tảng Công nghệ lõi (Tech Stack):\n"
     "- Backend (Xử lý logic & AI): Xây dựng theo kiến trúc hiện đại, linh hoạt với FastAPI (Python). Tích hợp Rule Engine nội bộ chuyên biệt cho việc tính toán điểm số KPI minh bạch theo các quy chế của cơ quan nhà nước.\n"
     "- Frontend (Giao diện người dùng): Sử dụng ReactJS và Vite, kết hợp thư viện UI Ant Design mang lại trải nghiệm người dùng tối giản, thân thiện nhưng vẫn đảm bảo sự chuyên nghiệp, chuẩn mực của ứng dụng chính phủ.\n"
     "- Hệ quản trị CSDL & Xử lý Dữ liệu lớn: Sử dụng PostgreSQL là CSDL chính, tích hợp pgvector cho các tác vụ nhúng vector và tìm kiếm AI. Đồng thời tích hợp cơ sở dữ liệu đồ thị KùzuDB (dạng embedded) để tăng tốc độ truy vấn các mối quan hệ phức tạp.\n"
     "- Mô hình AI & LLM: Hệ thống được thiết kế để có thể \"cắm rút\" (plug-and-play) linh hoạt với nhiều loại LLM khác nhau thông qua API (như GPT-4o-mini qua OpenRouter cho tính năng tiên tiến, hoặc các mô hình mã nguồn mở Llama-3, Qwen triển khai nội bộ để đảm bảo tuyệt đối an toàn thông tin).\n\n"
     "Phương thức triển khai (Deployment):\n"
     "- Hệ thống được đóng gói dưới dạng Ứng dụng Web (Web-based Application).\n"
     "- Có thể triển khai trên môi trường Đám mây Chính phủ (Government Cloud) hoặc triển khai nội bộ (On-premise) ngay tại các trung tâm dữ liệu của các cơ quan bộ, ban, ngành, đảm bảo tuân thủ các tiêu chuẩn bảo mật dữ liệu cấp quốc gia."),
    
    ("5. Đối tượng phục vụ:", 
     "Hệ thống được phân quyền tinh gọn, đáp ứng nhu cầu của 3 nhóm đối tượng chính trong Cơ quan Hành chính Nhà nước:\n"
     "- Ban Lãnh đạo cấp cao (Giám đốc, Phó giám đốc Sở / Lãnh đạo Bộ, Ngành): Đối tượng sử dụng chính của tính năng AI Copilot và Dashboard tổng quan. Cần nắm bắt thông tin nhanh chóng, giám sát tiến độ vĩ mô và ra quyết định chiến lược.\n"
     "- Cán bộ Quản lý cấp trung (Trưởng, Phó các phòng ban chuyên môn): Sử dụng hệ thống để phân bổ công việc, duyệt minh chứng do cấp dưới trình lên, theo dõi sát sao điểm số KPI của phòng ban mình để có các điều chỉnh kịp thời.\n"
     "- Công chức, Viên chức, Chuyên viên: Sử dụng để ghi nhận tiến độ công việc hằng ngày, nhận nhiệm vụ mới và tải lên các tài liệu minh chứng tự động mà không cần điền nhiều biểu mẫu phức tạp."),
    
    ("6. Kết quả mong đợi:", 
     "- Đột phá về Năng suất: Giải phóng hàng ngàn giờ làm việc mỗi năm cho cán bộ công chức khỏi các quy trình báo cáo, thủ tục giấy tờ mang tính thủ công. Tối ưu hóa 80-90% thời gian đánh giá minh chứng.\n"
     "- Nâng cao tính minh bạch và công bằng: Việc áp dụng Rule Engine và AI để đánh giá KPI sẽ dựa hoàn toàn trên dữ liệu thực tế (Data-driven), loại bỏ sự thiên vị hoặc cảm tính cá nhân, giúp kiến tạo một môi trường làm việc công bằng, thúc đẩy động lực cho cán bộ.\n"
     "- Quản trị rủi ro chủ động: Thay vì \"chữa cháy\", lãnh đạo có thể phòng ngừa các nguy cơ chậm trễ bằng hệ thống Risk Heatmap cảnh báo sớm, góp phần nâng cao năng lực phục vụ hành chính công.\n"
     "- Chuyển đổi số toàn diện: Giải pháp là bước đệm quan trọng trong công cuộc xây dựng Chính quyền số, Chính phủ điện tử, biến dữ liệu cứng nhắc thành tri thức phục vụ cho chỉ đạo, điều hành của nhà nước một cách thông minh, hiện đại.")
]

for title_text, content_text in sections:
    p_title = doc.add_paragraph()
    runner = p_title.add_run(title_text)
    runner.bold = True
    
    p_content = doc.add_paragraph()
    p_content.add_run(content_text)

doc.save('Ban_Mo_Ta_Y_Tuong_San_Pham_Chi_Tiet.docx')
print("Document saved successfully.")
