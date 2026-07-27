import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = docx.Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)

# Title
title = doc.add_paragraph('THÔNG TIN CHI TIẾT BẢN QUYỀN SẢN PHẨM')
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title.runs[0].bold = True
title.runs[0].font.size = Pt(15)

# Content sections
sections = [
    ("2. Chi tiết bản quyền của sản phẩm:", 
     "a) Các thành phần do nhóm tác giả hoàn toàn tự nghiên cứu, thiết kế và phát triển:\n"
     "- Kiến trúc hệ thống tổng thể: Thiết kế luồng nghiệp vụ và mô hình dữ liệu bám sát theo quy định về đánh giá, xếp loại chất lượng cán bộ, công chức, viên chức của Nhà nước.\n"
     "- Khối Logic nghiệp vụ (Rule Engine): Bộ quy tắc chấm điểm và đánh giá KPI được xây dựng riêng, có khả năng tự động tổng hợp, đối chiếu kết quả thực hiện nhiệm vụ với mục tiêu đề ra.\n"
     "- Hệ thống Trợ lý ảo (AI Copilot): Tự thiết kế hệ thống kịch bản (prompts) và luồng xử lý riêng để Trợ lý ảo hiểu và trả lời chính xác các nghiệp vụ, thuật ngữ hành chính công bằng tiếng Việt.\n"
     "- Hệ thống Phân tích minh chứng: Thuật toán tự động bóc tách thông tin từ các tài liệu, báo cáo (văn bản, PDF) do chuyên viên nộp lên để hỗ trợ công tác thẩm định tự động.\n"
     "- Bản đồ nhiệt (Risk Heatmap): Thuật toán phân tích mối liên kết giữa các phòng ban, cá nhân và đầu việc để phát hiện, cảnh báo sớm các điểm nghẽn, rủi ro chậm trễ.\n"
     "- Giao diện người dùng: Thiết kế Bảng điều khiển (Dashboard) quản trị trực quan, thân thiện, phù hợp với thói quen sử dụng của cán bộ nhà nước.\n\n"
     "b) Các thành phần kế thừa, sử dụng lại từ các nền tảng mã nguồn mở (có sẵn):\n"
     "- Lõi Xử lý Ngôn ngữ tự nhiên (LLM): Kế thừa các Mô hình ngôn ngữ lớn (Large Language Models) tiêu chuẩn để làm nền tảng trí tuệ nhân tạo cốt lõi cho việc phân tích văn bản và sinh báo cáo tự động.\n"
     "- Khối Lưu trữ dữ liệu lớn: Kế thừa các hệ quản trị cơ sở dữ liệu quan hệ, cơ sở dữ liệu véc-tơ (Vector DB) và đồ thị tri thức (Knowledge Graph) mã nguồn mở để lưu trữ an toàn thông tin nghiệp vụ và tối ưu hóa tìm kiếm.\n"
     "- Công cụ lập trình: Ứng dụng được phát triển dựa trên các nền tảng (framework) lập trình web tiêu chuẩn, mã nguồn mở, phổ biến và được cộng đồng quốc tế công nhận."),
    
    ("3. Danh sách các tài liệu tham khảo chứng minh bản quyền của sản phẩm:", 
     "- Kho mã nguồn (Source code) toàn bộ hệ thống do nhóm tác giả tự lập trình và lưu trữ.\n"
     "- Tài liệu thuyết minh giải pháp kỹ thuật, kiến trúc hệ thống đính kèm trong hồ sơ dự thi.\n"
     "- Phiên bản Demo chạy trực tuyến chứng minh đầy đủ các chức năng thực tế của sản phẩm.")
]

for title_text, content_text in sections:
    p_title = doc.add_paragraph()
    runner = p_title.add_run(title_text)
    runner.bold = True
    
    p_content = doc.add_paragraph()
    p_content.add_run(content_text)

doc.save('Chi_Tiet_Ban_Quyen_Du_Thi_Nha_Nuoc.docx')
print("Document saved successfully.")
